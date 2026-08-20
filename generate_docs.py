import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Create Documentation Directory
os.makedirs("Documentation", exist_ok=True)

doc = Document()

# Define global font settings
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Define Heading 1 style
h1_style = doc.styles['Heading 1']
h1_font = h1_style.font
h1_font.name = 'Times New Roman'
h1_font.size = Pt(16)
h1_font.bold = True
h1_font.color.rgb = None # Remove default blue color

# Helper functions
def add_justified_paragraph(text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    return p

def add_heading(text):
    h = doc.add_paragraph(text, style='Heading 1')
    return h

# Title Page Elements
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
t_run = title.add_run("POTHOLE SEVERITY CLASSIFIER & MAINTENANCE DECISION SUPPORT SYSTEM")
t_run.bold = True
t_run.font.size = Pt(16)

doc.add_paragraph("") # Spacing

sub1 = doc.add_paragraph()
sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub1_run = sub1.add_run("Final Documentation Report")
sub1_run.bold = True
sub1_run.font.size = Pt(14)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2_run = sub2.add_run("Project Report Submitted for Internship Project Submission")
sub2_run.bold = True
sub2_run.font.size = Pt(12)

doc.add_paragraph("")
doc.add_paragraph("")

add_justified_paragraph("Submitted By", bold=True)
add_justified_paragraph("Name: Shivam Subedar Patel")
add_justified_paragraph("Internship Role: Intern (AI/ML)")

doc.add_paragraph("")
add_justified_paragraph("Project Title: Pothole Severity Classifier & Maintenance Decision Support System", bold=True)

doc.add_paragraph("")
add_justified_paragraph("Technology Stack:", bold=True)
add_justified_paragraph("Python | Machine Learning | Scikit-Learn | Support Vector Machine | XGBoost | Pandas | NumPy | Streamlit | Plotly")

doc.add_paragraph("")
add_justified_paragraph("Project Description:", bold=True)
add_justified_paragraph("An end-to-end machine learning system developed to classify pothole severity from road inspection images, generate prioritization scores, and estimate repair costs based on scenario assumptions. The project uses machine learning models (Support Vector Machine) for severity prediction based on spatial geometry, computes an Expected Severity Priority Score, and features an interactive Streamlit dashboard with interactive Plotly metrics to visualize inspection hotspots and support municipal maintenance budget planning.")

doc.add_paragraph("")
add_justified_paragraph("Submitted To", bold=True)
add_justified_paragraph("Company/Organization Name: Data Vidwan")
add_justified_paragraph("Mentor/Manager Name: Meet Mistry")
add_justified_paragraph("Date of Submission: 20/08/2026")

doc.add_page_break()

# Main Content
add_heading("1. Project Objective")
add_justified_paragraph("The objective of this project is to develop a machine learning system that classifies pothole severity using historical road inspection imagery and spatial bounding box dimensions.")
add_justified_paragraph("The system predicts pothole severity (Minor, Medium, Major), identifies high-priority inspection targets based on a probability-weighted scoring framework, and provides a fully interactive UI. This enables municipalities to optimize maintenance budgets, plan repair schedules efficiently, and distribute resources without relying on historical financial receipts.")
add_justified_paragraph("The project includes:")
add_justified_paragraph("• Exploratory Data Analysis (EDA)")
add_justified_paragraph("• Data preprocessing and bounding box cleaning")
add_justified_paragraph("• Feature engineering (relative bounding dimensions, aspect ratios)")
add_justified_paragraph("• Multi-algorithm severity classification")
add_justified_paragraph("• Model comparison and hyperparameter tuning")
add_justified_paragraph("• Prioritization scoring (Critical/High/Medium/Low)")
add_justified_paragraph("• Scenario-based maintenance cost planning (USD/INR)")
add_justified_paragraph("• Interactive Streamlit dashboard")

doc.add_paragraph("")
add_heading("2. Requirement Coverage")
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Requirement'
hdr_cells[1].text = 'Status'
hdr_cells[2].text = 'Project File / Output'

reqs = [
    ("Severity Classification Model", "Completed", "models/final/final_model.joblib"),
    ("Feature Engineering", "Completed", "notebooks/02_feature_engineering_and_data_splitting.ipynb"),
    ("Cleaned Dataset", "Completed", "dataset/processed/pothole_dataset_cleaned.csv"),
    ("Priority Scoring System", "Completed", "src/priority_analysis.py"),
    ("Cost Estimation Model", "Completed", "src/cost_estimation.py"),
    ("EDA Report", "Completed", "reports/Final_Requirement_Audit.md"),
    ("Model Evaluation", "Completed", "reports/Municipal_Repair_Planning_Report.md"),
    ("Streamlit Dashboard", "Completed", "dashboard/app.py"),
    ("Project Documentation", "Completed", "Documentation/Final_Documentation_Report.docx")
]

for req, stat, out in reqs:
    row = table.add_row().cells
    row[0].text = req
    row[1].text = stat
    row[2].text = out
    # Set justification for table text
    for cell in row:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph("")
add_heading("3. Dataset")
add_justified_paragraph("The project uses a Pascal VOC image/object-detection dataset consisting of 717 images and 1,886 parsed pothole bounding boxes.")
add_justified_paragraph("Features:", bold=True)
add_justified_paragraph("• bbox_width, bbox_height, bbox_area (Spatial measurements)")
add_justified_paragraph("• rel_bbox_area (Relative dimension against original image size)")
add_justified_paragraph("• severity (Target Variable: minor_pothole, medium_pothole, major_pothole)")

doc.add_paragraph("")
add_heading("4. Work Completed")
add_justified_paragraph("Completed an end-to-end AI/ML workflow including:")
add_justified_paragraph("• Loaded and inspected raw XML annotations.")
add_justified_paragraph("• Performed exploratory data analysis and spatial distribution mapping.")
add_justified_paragraph("• Extracted spatial bounding box geometry into analytical features.")
add_justified_paragraph("• Trained 5 distinct machine learning algorithms (Random Forest, XGBoost, LightGBM, Gradient Boosting, Support Vector Machine).")
add_justified_paragraph("• Compared model performance and evaluated metrics (Accuracy, Precision, Recall, Macro F1).")
add_justified_paragraph("• Selected Support Vector Machine for final deployment.")
add_justified_paragraph("• Implemented a continuous priority scoring formula based on class probabilities.")
add_justified_paragraph("• Developed a fully functional, multi-page Streamlit dashboard with premium glassmorphism aesthetics.")
add_justified_paragraph("• Exported processed datasets, detailed markdown reports, and trained models.")

doc.add_paragraph("")
add_heading("5. Methodology")
add_justified_paragraph("5.1 Exploratory Data Analysis", bold=True)
add_justified_paragraph("Performed analysis on:")
add_justified_paragraph("• Target class distribution (minor, medium, major)")
add_justified_paragraph("• Dimensional correlation mapping")
add_justified_paragraph("• Image bounding box anomalies")

doc.add_paragraph("")
add_justified_paragraph("5.2 Feature Engineering", bold=True)
add_justified_paragraph("Created additional analytical features:")
add_justified_paragraph("• Relative bounding box geometries (rel_bbox_width, rel_bbox_height, rel_bbox_area)")
add_justified_paragraph("• Aspect Ratios (bbox_aspect_ratio)")
add_justified_paragraph("• Logarithmic transformations (log_bbox_area)")

doc.add_paragraph("")
add_heading("6. Model Development")
add_justified_paragraph("Multiple classification models were trained using robust Leak-Safe GroupKFold cross-validation:")
add_justified_paragraph("• Support Vector Machine")
add_justified_paragraph("• Gradient Boosting")
add_justified_paragraph("• LightGBM")
add_justified_paragraph("• Random Forest")
add_justified_paragraph("• XGBoost")
add_justified_paragraph("The models were evaluated using Accuracy and Macro F1 Score to account for class imbalances.")

doc.add_paragraph("")
add_heading("7. Model Results")
add_justified_paragraph("Baseline validation results prior to final tuning:")
res_table = doc.add_table(rows=1, cols=3)
res_table.style = 'Table Grid'
h = res_table.rows[0].cells
h[0].text = 'Model'
h[1].text = 'Macro F1 (%)'
h[2].text = 'Accuracy (%)'
results = [
    ("Support Vector Machine", "62.46", "64.03"),
    ("Gradient Boosting", "58.26", "59.74"),
    ("LightGBM", "56.66", "58.42"),
    ("Random Forest", "56.20", "58.09"),
    ("XGBoost", "54.71", "57.10")
]
for r1, r2, r3 in results:
    row = res_table.add_row().cells
    row[0].text = r1
    row[1].text = r2
    row[2].text = r3

doc.add_paragraph("")
add_heading("8. Selected Model")
add_justified_paragraph("Support Vector Machine (RBF Kernel)", bold=True)
add_justified_paragraph("Reason for selection:", bold=True)
add_justified_paragraph("The SVM significantly outperformed all tree-based ensembles on the validation groups for this specific geometric dataset, providing the highest Macro F1 score and the best robustness against overfitting.")

doc.add_paragraph("")
add_heading("9. Streamlit Application")
add_justified_paragraph("The Streamlit application is implemented in:")
add_justified_paragraph("dashboard/app.py")
add_justified_paragraph("The dashboard contains:")
add_justified_paragraph("• Executive Overview: Key Performance Indicators and final dataset previews.")
add_justified_paragraph("• Severity Analysis: Visual distributions and probability heatmaps.")
add_justified_paragraph("• Priority Framework: Histograms mapped to decision-support tiers (Critical to Low).")
add_justified_paragraph("• Hotspot Identification: Image-level aggregation to prioritize inspection locales.")
add_justified_paragraph("• Scenario-Based Cost Planning: Interactive USD/INR estimated budget assumptions.")

doc.add_paragraph("")
add_heading("10. Key Findings")
add_justified_paragraph("• Extracted bounding box area possesses a direct correlation to labeled severity.")
add_justified_paragraph("• Tree-based models severely overfit spatial coordinates; scaled distance-based SVMs provide superior generalization.")
add_justified_paragraph("• Utilizing prediction probabilities generates a highly effective continuous triage score.")

doc.add_paragraph("")
add_heading("11. Recommendations")
add_justified_paragraph("• Incorporate external hardware GPS modules during image capture for geographic mapping.")
add_justified_paragraph("• Introduce LiDAR or depth-sensing arrays to compute true volumetric damage instead of 2D bounding boxes.")
add_justified_paragraph("• Cross-reference municipal receipts to build a supervised cost regression model rather than relying on heuristic scenarios.")

doc.add_paragraph("")
add_heading("12. Achievements")
add_justified_paragraph("Successfully completed:")
add_justified_paragraph("• Data preprocessing and bounding box validation.")
add_justified_paragraph("• Machine learning model development and comprehensive evaluation (5 models).")
add_justified_paragraph("• Robust 4-tier decision-support priority engine.")
add_justified_paragraph("• End-to-End multi-tabbed Streamlit Dashboard deployment with premium Plotly aesthetics.")

doc.add_paragraph("")
add_heading("13. Limitations")
add_justified_paragraph("• Dataset lacks exogenous attributes such as Traffic Volume (AADT), geographic coordinates, historical repair costs, and road condition variables (IRI/PCI).")

doc.add_paragraph("")
add_heading("14. Final Submission Contents")
add_justified_paragraph("• Documentation Report (.docx)")
add_justified_paragraph("• Machine Learning Models (.joblib)")
add_justified_paragraph("• Streamlit Dashboard (dashboard/app.py)")
add_justified_paragraph("• Processed Datasets (outputs/)")
add_justified_paragraph("• Markdown Reports (reports/)")
add_justified_paragraph("• Jupyter Notebooks (notebooks/)")

doc.add_paragraph("")
add_heading("15. How to Run")
add_justified_paragraph("pip install -r requirements.txt")
add_justified_paragraph("python -m streamlit run dashboard/app.py")

doc.add_paragraph("")
add_heading("16. Final Conclusion")
add_justified_paragraph("The Pothole Severity Classifier & Decision Support System successfully delivers an end-to-end machine learning solution for damage triage, severity forecasting, and transparent financial planning.")
add_justified_paragraph("The system combines spatial AI with robust real-time heuristic modeling to help municipal managers optimize maintenance deployment and make data-driven decisions through a fully interactive, premium Streamlit dashboard.")

doc.save("Documentation/Final_Documentation_Report.docx")
